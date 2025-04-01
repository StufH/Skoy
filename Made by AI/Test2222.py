# -*- coding: utf-8 -*-
# VIKTIG: Denne koden krever at du har installert pandas, matplotlib OG python-docx
# Kjør: pip install pandas matplotlib python-docx
# Koden vil sannsynligvis feile i dette miljøet pga. manglende python-docx.
# Selv om den kjører, kan ikke den resulterende .docx-filen leveres.

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import os
import time # For å lage unike filnavn
import sys # For å kunne avslutte

# --- Importer for Word-dokument ---
try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("="*50)
    print("FEIL: Biblioteket 'python-docx' er ikke installert.")
    print("Kjør 'pip install python-docx' på din egen maskin for å bruke denne funksjonaliteten.")
    print("Denne koden kan ikke lage Word-filen i dette miljøet.")
    print("="*50)
    CAN_RUN_DOCX = False
else:
    CAN_RUN_DOCX = True


# Ensure matplotlib uses a non-interactive backend
import matplotlib
matplotlib.use('Agg')

# ===============================================
# 1. DEFINER FUNKSJON FOR PLOTTING / DATAANALYSE
#    (Identisk med forrige versjon - genererer alle 7 bilder)
# ===============================================
def generate_plots_and_get_filenames(initial_data):
    """Kjører dataanalyse, genererer ALLE 7 PNG-filer og returnerer filnavnene."""
    print("Starter dataanalyse og plotting...")
    plot_filenames = {}
    timestamp_internal = int(time.time())
    base_names = {
        'all_spm': 'graf_alle_grupper_alle_spm', 'g1_spm': 'graf_gruppe_16-21',
        'g2_spm': 'graf_gruppe_30-50', 'spm2_comp': 'graf_alle_grupper_spm2',
        'spm3_comp': 'graf_alle_grupper_spm3', 'emo_match': 'graf_rett_emosjon_prosent',
        'emo_table': 'tabell_rett_emosjon_status'
    }
    png_filenames = {key: f"{name}_{timestamp_internal}.png" for key, name in base_names.items()}
    EXPECTED_COLUMNS = ['AgeGroup', 'ParticipantID', 'ImageType', 'ImageName', 'Spm1', 'Spm2', 'Spm3']
    correct_emotion_map = {
        'kjedsomhet': ['kjedsomhet', 'kjedelig'], 'nostalgi': ['nostalgi', 'nostalgisk'],
        'sinne': ['sinne'], 'frykt': ['frykt', 'krig/frykt'], 'avsky': ['avsky', 'kvalme'],
        'glede': ['glede', 'glad', 'håp'], 'overraskelse': ['overraskelse', 'overrasket', 'overveldende'],
        'tristhet': ['tristhet', 'trist']
    }
    def is_emotion_correct(row):
        theme = str(row['ImageName']).lower().strip()
        perceived = str(row['Spm1']).lower().strip()
        normalized_map = {k.lower(): [v.lower() for v in val] for k, val in correct_emotion_map.items()}
        is_correct = False
        if theme in normalized_map: is_correct = perceived in normalized_map[theme]
        return 1 if is_correct else 0
    df = pd.DataFrame(initial_data)
    if 'CorrectEmotion' not in df.columns: df['CorrectEmotion'] = df.apply(is_emotion_correct, axis=1)
    try:
        if df.empty: raise ValueError("DataFrame er tom.")
        cols_to_check = EXPECTED_COLUMNS
        missing_cols = [col for col in cols_to_check if col not in df.columns]
        if missing_cols: raise ValueError(f"Mangler kolonner: {missing_cols}")
        df['Spm2'] = pd.to_numeric(df['Spm2'], errors='coerce')
        df['Spm3'] = pd.to_numeric(df['Spm3'], errors='coerce')
        df['Spm1'] = df['Spm1'].astype(str); df['ImageName'] = df['ImageName'].astype(str)
        df['AgeGroup'] = df['AgeGroup'].astype(str); df['ImageType'] = df['ImageType'].astype(str)
        df['ParticipantID'] = df['ParticipantID'].astype(str)
        analysis_cols = ['AgeGroup', 'ImageType', 'Spm2', 'Spm3', 'CorrectEmotion', 'ImageName', 'Spm1']
        df.dropna(subset=analysis_cols, inplace=True)
        if df.empty: raise ValueError("Ingen gyldige rader etter fjerning av NaN.")
    except Exception as e: print(f"FEIL datavalidering: {e}"); return None
    plot_data_spm2, plot_data_spm3, plot_data_emo_match = pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
    age_groups_found = []; averages_dict = {}; df_filtered = pd.DataFrame()
    try:
        df_filtered = df[df['ImageType'].isin(['Menneske', 'KI'])].copy()
        if df_filtered.empty: raise ValueError("Ingen rader med ImageType 'Menneske' eller 'KI'.")
        average_scores_calc = df_filtered.groupby(['AgeGroup', 'ImageType'])[['Spm2', 'Spm3', 'CorrectEmotion']].mean()
        for index, row in average_scores_calc.iterrows():
             age_group, image_type = index
             if age_group not in averages_dict: averages_dict[age_group] = {}
             averages_dict[age_group][image_type] = row.to_dict()
        average_scores_unstacked = average_scores_calc.unstack()
        if isinstance(average_scores_unstacked.columns, pd.MultiIndex):
            average_scores_unstacked.columns = ['_'.join(map(str, col)).strip() for col in average_scores_unstacked.columns.values]
        average_scores_unstacked = average_scores_unstacked.fillna(0)
        age_groups_found = average_scores_unstacked.index.unique().tolist()
        if not age_groups_found: raise ValueError("Ingen aldersgrupper funnet.")
        required_cols = ['Spm2_Menneske', 'Spm2_KI', 'Spm3_Menneske', 'Spm3_KI', 'CorrectEmotion_Menneske', 'CorrectEmotion_KI']
        for col in required_cols:
            if col not in average_scores_unstacked.columns: average_scores_unstacked[col] = 0.0
        plot_data_spm2 = average_scores_unstacked[['Spm2_Menneske', 'Spm2_KI']].loc[age_groups_found]
        plot_data_spm3 = average_scores_unstacked[['Spm3_Menneske', 'Spm3_KI']].loc[age_groups_found]
        plot_data_emo_match = average_scores_unstacked[['CorrectEmotion_Menneske', 'CorrectEmotion_KI']].loc[age_groups_found]
        print("Gjennomsnitt beregnet.")
    except Exception as e: print(f"FEIL beregning: {e}"); return None
    if plot_data_spm2.empty or plot_data_spm3.empty or plot_data_emo_match.empty or not age_groups_found:
        print("Kan ikke plotte, mangler data."); return None
    image_types_plot = ['Menneskeskapt', 'KI-generert']; bar_width = 0.35
    index_groups = np.arange(len(age_groups_found)); plot_counter = 0
    # Plot 1 ('all_spm')
    try:
        fig1, axes1 = plt.subplots(1, 2, figsize=(14, 6)); fig1.suptitle('Gj.sn. Respons & Behagelighet: Sammenligning', fontsize=16, y=1.03)
        rects1 = axes1[0].bar(index_groups - bar_width/2, plot_data_spm2['Spm2_Menneske'], bar_width, label='Menneskeskapt', color='skyblue')
        rects2 = axes1[0].bar(index_groups + bar_width/2, plot_data_spm2['Spm2_KI'], bar_width, label='KI-generert', color='lightcoral')
        axes1[0].set_ylabel('Gj.sn. score (1-10)'); axes1[0].set_title('Intensitet (Spm 2)'); axes1[0].set_xticks(index_groups); axes1[0].set_xticklabels(age_groups_found)
        axes1[0].set_ylim(0, 11); axes1[0].legend(); axes1[0].grid(axis='y', linestyle='--'); axes1[0].bar_label(rects1, fmt='%.2f', fontsize=9); axes1[0].bar_label(rects2, fmt='%.2f', fontsize=9)
        rects3 = axes1[1].bar(index_groups - bar_width/2, plot_data_spm3['Spm3_Menneske'], bar_width, label='Menneskeskapt', color='skyblue')
        rects4 = axes1[1].bar(index_groups + bar_width/2, plot_data_spm3['Spm3_KI'], bar_width, label='KI-generert', color='lightcoral')
        axes1[1].set_ylabel('Gj.sn. score (1-10)'); axes1[1].set_title('Behagelighet (Spm 3)'); axes1[1].set_xticks(index_groups); axes1[1].set_xticklabels(age_groups_found)
        axes1[1].set_ylim(0, 11); axes1[1].legend(); axes1[1].grid(axis='y', linestyle='--'); axes1[1].bar_label(rects3, fmt='%.2f', fontsize=9); axes1[1].bar_label(rects4, fmt='%.2f', fontsize=9)
        plt.tight_layout(rect=[0, 0.03, 1, 0.95]); plt.savefig(png_filenames['all_spm']); plot_filenames['all_spm'] = png_filenames['all_spm']; print(f"Lagret: {png_filenames['all_spm']}")
        plt.close(fig1); plot_counter += 1
    except Exception as e: print(f"FEIL plot 1: {e}")
    # Plot 2 ('g1_spm')
    age_group_1 = "16-21"
    if age_group_1 in averages_dict:
        try:
            fig2, axes2 = plt.subplots(1, 2, figsize=(12, 5), sharey=True); fig2.suptitle(f'Gjennomsnittsresponser - Aldersgruppe: {age_group_1}', fontsize=16)
            avg_data = averages_dict[age_group_1]; spm2_m = avg_data.get('Menneske', {}).get('Spm2', 0); spm2_k = avg_data.get('KI', {}).get('Spm2', 0); spm3_m = avg_data.get('Menneske', {}).get('Spm3', 0); spm3_k = avg_data.get('KI', {}).get('Spm3', 0)
            bars1 = axes2[0].bar(image_types_plot, [spm2_m, spm2_k], color=['skyblue', 'lightcoral']); axes2[0].set_title('Intensitet (Spm 2)'); axes2[0].set_ylabel('Gj.sn. score (1-10)'); axes2[0].set_ylim(0, 11); axes2[0].grid(axis='y', linestyle='--'); axes2[0].bar_label(bars1, fmt='%.2f', fontsize=9)
            bars2 = axes2[1].bar(image_types_plot, [spm3_m, spm3_k], color=['skyblue', 'lightcoral']); axes2[1].set_title('Behagelighet (Spm 3)'); axes2[1].grid(axis='y', linestyle='--'); axes2[1].bar_label(bars2, fmt='%.2f', fontsize=9)
            plt.tight_layout(rect=[0, 0.03, 1, 0.95]); plt.savefig(png_filenames['g1_spm']); plot_filenames['g1_spm'] = png_filenames['g1_spm']; print(f"Lagret: {png_filenames['g1_spm']}")
            plt.close(fig2); plot_counter += 1
        except Exception as e: print(f"FEIL plot 2 ({age_group_1}): {e}")
    else: print(f"Ingen data for plot 2 ({age_group_1}).")
    # Plot 3 ('g2_spm')
    age_group_2 = "30-50"
    if age_group_2 in averages_dict:
        try:
            fig3, axes3 = plt.subplots(1, 2, figsize=(12, 5), sharey=True); fig3.suptitle(f'Gjennomsnittsresponser - Aldersgruppe: {age_group_2}', fontsize=16)
            avg_data = averages_dict[age_group_2]; spm2_m = avg_data.get('Menneske', {}).get('Spm2', 0); spm2_k = avg_data.get('KI', {}).get('Spm2', 0); spm3_m = avg_data.get('Menneske', {}).get('Spm3', 0); spm3_k = avg_data.get('KI', {}).get('Spm3', 0)
            bars1 = axes3[0].bar(image_types_plot, [spm2_m, spm2_k], color=['skyblue', 'lightcoral']); axes3[0].set_title('Intensitet (Spm 2)'); axes3[0].set_ylabel('Gj.sn. score (1-10)'); axes3[0].set_ylim(0, 11); axes3[0].grid(axis='y', linestyle='--'); axes3[0].bar_label(bars1, fmt='%.2f', fontsize=9)
            bars2 = axes3[1].bar(image_types_plot, [spm3_m, spm3_k], color=['skyblue', 'lightcoral']); axes3[1].set_title('Behagelighet (Spm 3)'); axes3[1].grid(axis='y', linestyle='--'); axes3[1].bar_label(bars2, fmt='%.2f', fontsize=9)
            plt.tight_layout(rect=[0, 0.03, 1, 0.95]); plt.savefig(png_filenames['g2_spm']); plot_filenames['g2_spm'] = png_filenames['g2_spm']; print(f"Lagret: {png_filenames['g2_spm']}")
            plt.close(fig3); plot_counter += 1
        except Exception as e: print(f"FEIL plot 3 ({age_group_2}): {e}")
    else: print(f"Ingen data for plot 3 ({age_group_2}).")
    # Plot 4 ('spm2_comp')
    try:
        fig4, ax4 = plt.subplots(1, 1, figsize=(7, 6)); fig4.suptitle('Gj.sn. Respons/Intensitet (Spm 2): Sammenligning', fontsize=14, y=1.0)
        rects1 = ax4.bar(index_groups - bar_width/2, plot_data_spm2['Spm2_Menneske'], bar_width, label='Menneskeskapt', color='skyblue')
        rects2 = ax4.bar(index_groups + bar_width/2, plot_data_spm2['Spm2_KI'], bar_width, label='KI-generert', color='lightcoral')
        ax4.set_ylabel('Gj.sn. score (1-10)'); ax4.set_title('Intensitet (Spm 2)'); ax4.set_xticks(index_groups); ax4.set_xticklabels(age_groups_found); ax4.set_ylim(0, 11); ax4.legend(); ax4.grid(axis='y', linestyle='--')
        ax4.bar_label(rects1, fmt='%.2f', fontsize=9); ax4.bar_label(rects2, fmt='%.2f', fontsize=9)
        plt.tight_layout(rect=[0, 0.03, 1, 0.95]); plt.savefig(png_filenames['spm2_comp']); plot_filenames['spm2_comp'] = png_filenames['spm2_comp']; print(f"Lagret: {png_filenames['spm2_comp']}")
        plt.close(fig4); plot_counter += 1
    except Exception as e: print(f"FEIL plot 4: {e}")
    # Plot 5 ('spm3_comp')
    try:
        fig5, ax5 = plt.subplots(1, 1, figsize=(7, 6)); fig5.suptitle('Gj.sn. Behagelighet (Spm 3): Sammenligning', fontsize=14, y=1.0)
        rects3 = ax5.bar(index_groups - bar_width/2, plot_data_spm3['Spm3_Menneske'], bar_width, label='Menneskeskapt', color='skyblue')
        rects4 = ax5.bar(index_groups + bar_width/2, plot_data_spm3['Spm3_KI'], bar_width, label='KI-generert', color='lightcoral')
        ax5.set_ylabel('Gj.sn. score (1-10)'); ax5.set_title('Behagelighet (Spm 3)'); ax5.set_xticks(index_groups); ax5.set_xticklabels(age_groups_found); ax5.set_ylim(0, 11); ax5.legend(); ax5.grid(axis='y', linestyle='--')
        ax5.bar_label(rects3, fmt='%.2f', fontsize=9); ax5.bar_label(rects4, fmt='%.2f', fontsize=9)
        plt.tight_layout(rect=[0, 0.03, 1, 0.95]); plt.savefig(png_filenames['spm3_comp']); plot_filenames['spm3_comp'] = png_filenames['spm3_comp']; print(f"Lagret: {png_filenames['spm3_comp']}")
        plt.close(fig5); plot_counter += 1
    except Exception as e: print(f"FEIL plot 5: {e}")
    # Plot 6 ('emo_match')
    try:
        fig6, ax6 = plt.subplots(1, 1, figsize=(7, 6)); fig6.suptitle('Andel "Rett" Emosjon (Spm 1 vs Bildetema)', fontsize=14, y=1.0)
        correct_emo_m = plot_data_emo_match['CorrectEmotion_Menneske'] * 100; correct_emo_k = plot_data_emo_match['CorrectEmotion_KI'] * 100
        rects5 = ax6.bar(index_groups - bar_width/2, correct_emo_m, bar_width, label='Menneskeskapt', color='skyblue')
        rects6 = ax6.bar(index_groups + bar_width/2, correct_emo_k, bar_width, label='KI-generert', color='lightcoral')
        ax6.set_ylabel('Prosentandel (%)'); ax6.set_title('"Rett" Emosjon (Spm 1)'); ax6.set_xticks(index_groups); ax6.set_xticklabels(age_groups_found); ax6.set_ylim(0, 105); ax6.legend(); ax6.grid(axis='y', linestyle='--')
        ax6.bar_label(rects5, fmt='%.0f%%', fontsize=9); ax6.bar_label(rects6, fmt='%.0f%%', fontsize=9)
        plt.tight_layout(rect=[0, 0.03, 1, 0.95]); plt.savefig(png_filenames['emo_match']); plot_filenames['emo_match'] = png_filenames['emo_match']; print(f"Lagret: {png_filenames['emo_match']}")
        plt.close(fig6); plot_counter += 1
    except Exception as e: print(f"FEIL plot 6: {e}")
    # Plot 7 ('emo_table')
    try:
        if not df_filtered.empty:
            table_data = df_filtered[['AgeGroup', 'ParticipantID', 'ImageName', 'ImageType', 'Spm1', 'CorrectEmotion']].copy()
            table_data['Status'] = table_data['CorrectEmotion'].apply(lambda x: '✓' if x==1 else '✗')
            table_data.sort_values(by=['AgeGroup','ParticipantID','ImageType'], inplace=True)
            table_display_data = table_data[table_data['ImageType'].isin(['Menneske','KI'])]
            table_display_data = table_display_data[['AgeGroup', 'ParticipantID', 'ImageName', 'ImageType', 'Spm1', 'Status']]
            col_labels = ['Aldersgr.', 'Deltaker', 'Bildetema', 'Bildetype', 'Opplevd (Spm1)', 'Rett?']
            row_height = 0.3; fig_height = max(3, len(table_display_data) * row_height)
            fig_table, ax_table = plt.subplots(figsize=(8, fig_height))
            ax_table.axis('tight'); ax_table.axis('off')
            if not table_display_data.empty:
                 the_table = ax_table.table(cellText=table_display_data.values, colLabels=col_labels, loc='center', cellLoc='left')
                 the_table.auto_set_font_size(False); the_table.set_fontsize(8); the_table.scale(1, 1)
            else: ax_table.text(0.5, 0.5, "Ingen data", ha='center', va='center')
            plt.title("Opplevd Emosjon Status", fontsize=10); plt.tight_layout(pad=0.2)
            plt.savefig(png_filenames['emo_table']); plot_filenames['emo_table'] = png_filenames['emo_table']; print(f"Lagret: {png_filenames['emo_table']}")
            plt.close(fig_table); plot_counter += 1
        else: print("FEIL plot 7: df_filtered er tom.")
    except Exception as e: print(f"FEIL plot 7: {e}")

    print(f"{plot_counter} bilder generert.")
    return plot_filenames if plot_counter >= 7 else None

# ===============================================
# 2. DEFINER RAPPORTTEKSTEN (FULLSTENDIG VERSJON)
# ===============================================
def get_full_report_text_complete():
    """Returnerer den komplette rapportteksten uten forkortelser."""
    report_content = [
        ('title', 'Psykologisk Eksperiment: Emosjonell Respons til KI-generert vs. Menneskeskapt Kunst'),
        ('p_center_bold', f"Dato: {time.strftime('%d. %B %Y')}"),
        ('h1', 'Innholdsfortegnelse'),
        ('toc', '1.0 Introduksjon .......................................................................................................................... X'),
        ('toc_sub', '1.1 Grunnlaget for forsøket .................................................................................................. X'),
        ('toc_sub', '1.2 Teoretisk rammeverk ...................................................................................................... X'),
        ('toc_ssub', '1.2.1 Persepsjon og kognisjon ....................................................................................... X'),
        ('toc_ssub', '1.2.2 Emosjoner ............................................................................................................. X'),
        ('toc_ssub', '1.2.3 Hvorfor får vi en emosjonell respons til kunst? ................................................... X'),
        ('toc_ssub', '1.2.4 Den uhyggelige dal (Uncanny Valley) .................................................................. X'),
        ('toc_ssub', '1.2.5 Ytterligere teoretiske perspektiver ................................................................... X'),
        ('toc_sub', '1.3 Metode ............................................................................................................................ X'),
        ('toc_sub', '1.4 Variabler .......................................................................................................................... X'),
        ('toc_ssub', '1.4.1 Uavhengige variabler ........................................................................................... X'),
        ('toc_ssub', '1.4.2 Avhengige variabler ............................................................................................ X'),
        ('toc_sub', '1.5 Hypoteser ........................................................................................................................ X'),
        ('toc', '2.0 Forsøksdesign ...................................................................................................................... X'),
        ('toc_sub', '2.1 Utvalg ............................................................................................................................. X'),
        ('toc_sub', '2.2 Valg av bilder ................................................................................................................. X'),
        ('toc_sub', '2.3 Eksperimentets gjennomføring ........................................................................................ X'),
        ('toc_sub', '2.4 Innsamling av data .......................................................................................................... X'),
        ('toc', '3.0 Resultater og Dataanalyse ................................................................................................... X'),
        ('toc_sub', '3.1 Kvantitative resultater (Spm 2 & Spm 3) ...................................................................... X'),
        ('toc_sub', '3.2 Kvalitative resultater (Spm 1 - "Rett Emosjon") .......................................................... X'),
        ('toc', '4.0 Validitet og reliabilitet ....................................................................................................... X'),
        ('toc_sub', '4.1 Styrker og svakheter ..................................................................................................... X'),
        ('toc_ssub', '4.1.1 Styrker ................................................................................................................ X'),
        ('toc_ssub', '4.1.2 Svakheter ............................................................................................................ X'),
        ('toc_sub', '4.2 Drøfting ........................................................................................................................ X'),
        ('toc', '5.0 Konklusjon ....................................................................................................................... X'),
        ('toc', '6.0 Referanser ........................................................................................................................ X'),
        ('toc', '7.0 Appendix .......................................................................................................................... X'),
        ('toc_sub', '7.1 Genererte Grafer ........................................................................................................... X'),
        ('toc_sub', '7.2 Tabell: Opplevd Emosjon Status ................................................................................... X'),
        ('toc_sub', '7.3 Definisjon av "Rett Emosjon" ...................................................................................... X'),
        ('page_break', ''),

        ('h1', '1.0 Introduksjon'), ('h2', '1.1 Grunnlaget for forsøket'),
        ('p', 'Kunstig intelligens (KI) har de siste årene demonstrert en stadig økende kapasitet til å skape fotorealistiske og kunstnerisk overbevisende bilder. Noen av disse verkene har blitt solgt for betydelige summer, noen ganger uten at kjøperne var klar over at opphavet ikke var menneskelig. Denne utviklingen reiser interessante spørsmål om hvordan vi som mennesker responderer på kunst skapt av maskiner kontra kunst skapt av mennesker. Derfor ønsket vi å undersøke om menneskets emosjonelle responser varierer systematisk når de observerer KI-genererte bilder sammenlignet med menneskeskapte bilder. Videre ville vi utforske om det eksisterer en forskjell i disse responsene mellom ulike aldersgrupper, spesifikt en yngre gruppe (16-21 år) antatt å være mer eksponert for digital teknologi, og en eldre gruppe (30-50 år).'),
        ('p', 'Dette psykologiske eksperimentet er relevant fordi det kan gi innsikt i hvordan KI-generert innhold påvirker oss på et fundamentalt emosjonelt nivå. Resultatene kan bidra til diskusjonen om autentisitet og verdi i kunstverdenen, og belyse behovet for eventuelle retningslinjer eller verktøy for å identifisere og håndtere potensielle forfalskninger eller misrepresentasjoner innen moderne kunst ved hjelp av kunstig intelligens.'),
        ('h2', '1.2 Teoretisk rammeverk'),
        ('p', 'For å forstå og tolke resultatene av eksperimentet, støtter vi oss på teorier innen persepsjon, kognisjon, emosjonspsykologi og menneske-maskin-interaksjon.'),
        ('h3', '1.2.1 Persepsjon og kognisjon'),
        ('p', 'Persepsjon er prosessen der hjernen organiserer og tolker sanseinformasjon (stimuli) for å skape en meningsfull forståelse av omverdenen. Når vi ser på et bilde, involverer dette komplekse kognitive prosesser som mønstergjenkjenning, farge- og formoppfatning, og aktivering av assosiasjoner og minner. Kognisjon omfatter de mentale prosessene knyttet til tenkning, læring, hukommelse og problemløsning, som alle spiller inn i hvordan vi tolker og vurderer et kunstverk.'),
        ('h3', '1.2.2 Emosjoner'),
        ('p', 'For å forstå de emosjonelle responsene som undersøkes, er det viktig å definere hva emosjoner innebærer. Emosjoner (følelser) er komplekse psykologiske tilstander som ofte beskrives med tre komponenter:'),
        ('list_num', '1.  Subjektiv opplevelse: Den bevisste følelsen individet opplever (f.eks. glede, sinne, tristhet). Dette er den komponenten vi primært måler gjennom selvrapportering i dette forsøket.'),
        ('list_num', '2.  Fysiologisk respons: Kroppslige endringer som aktiveres av emosjonen, styrt av det autonome nervesystemet (f.eks. endringer i hjerterate, svetteproduksjon, muskelspenninger).'),
        ('list_num', '3.  Atferdsmessig (ekspressiv) respons: Ytre handlinger eller uttrykk som følger av emosjonen (f.eks. ansiktsuttrykk, kroppsspråk, fluktrespons). (Svartdal, 2025)'),
        ('p', 'Paul Ekman (Dr. Paul Ekman, u.å.; «Emotion Classification», 2024) er kjent for sitt arbeid med å identifisere grunnleggende, universelle emosjoner som glede, tristhet, sinne, frykt, avsky og overraskelse. Han argumenterte for at disse er medfødte og gjenkjennelige på tvers av kulturer. Selv om kunst kan fremkalle disse grunnemosjonene, kan den også vekke mer sammensatte følelser som ærefrykt, undring, eller, som inkludert i dette studiet, nostalgi og kjedsomhet, som går utover Ekmans opprinnelige liste.'),
        ('h3', '1.2.3 Hvorfor får vi en emosjonell respons til kunst?'),
        ('p', 'Fenomenet at kunst kan vekke sterke følelser er velkjent, fra det angstfylte uttrykket i Edvard Munks «Skrik» til den ømme intimiteten i Gustav Klimts «Kysset». Eldre teorier foreslo at emosjonell respons på kunst primært skyldtes enkle stimulus-respons-mekanismer knyttet til farge, form og komposisjon (Silvia, 2005). Nyere forskning, derimot, peker på mer komplekse nevrologiske og kognitive faktorer.'),
        ('p', 'Mønstergjenkjenning er en fundamental kognitiv evne der hjernen sammenligner innkommende stimuli med lagrede minner og prototyper (Yin, 2008). Denne evnen, essensiell for overlevelse (identifisere farer, gjenkjenne mat, tolke sosiale signaler), aktiveres også når vi ser på kunst. Vi søker ubevisst etter kjente mønstre, former og fargekombinasjoner som kan vekke assosiasjoner og minner, og dermed følelser (Christensen, 2014).'),
        ('p', 'Empati og speilnevroner spiller også en sentral rolle. Speilnevroner er nerveceller som aktiveres både når vi utfører en handling *og* når vi observerer den samme handlingen utført av andre («Speilnevroner», 2024). Det antas at disse nevronene er involvert i vår evne til å forstå andres intensjoner og følelser (empati) (Christensen, 2014). Når vi ser et kunstverk som portretterer en figur som uttrykker en sterk følelse (f.eks. sorg, glede), kan våre speilnevroner aktiveres, noe som gir en følelse av emosjonell resonans med det avbildede.'),
        ('h3', '1.2.4 Den uhyggelige dal (Uncanny Valley)'),
        ('p', '«Den uhyggelige dal» er en hypotese innen estetikk og robotikk, opprinnelig formulert av Masahiro Mori («Den uhyggelige dal», 2024; Derfor er robotter og klovne så uhyggelige, 2020). Teorien postulerer at når en kunstig representasjon (som en robot eller animert figur) nærmer seg et menneskelig utseende, øker vår positive affinitet for den – men bare opp til et visst punkt. Hvis likheten blir *nesten* perfekt, men med subtile avvik eller "feil", faller vår affinitet dramatisk og kan erstattes av en følelse av ubehag, avsky eller uhygge. Først når representasjonen blir *svært* lik et menneske (vanskelig å skille fra ekte), kan affiniteten øke igjen.'),
        ('p', '(Her skulle Bilde 1 [Grafen av Uncanny Valley, MacDorman, 2005] normalt settes inn manuelt)'),
        ('p', 'Interessant nok tyder forskning på at denne effekten ikke er unik for mennesker. En studie viste at aper (makaker) foretrakk å se på bilder av ekte apeansikter eller klart urealistiske syntetiske ansikter fremfor svært realistiske, men ikke perfekte, syntetiske apeansikter (Steckenfinger & Ghazanfar, 2009; Oct. 13 et al., u.å.). Dette funnet antyder at fenomenet kan ha dype evolusjonære røtter, muligens knyttet til mekanismer for å gjenkjenne og unngå syke individer, artsfrender med genetiske defekter, eller døde kropper.'),
        ('p', 'KI-genererte bilder, spesielt de som forsøker å etterligne virkeligheten (som portretter), kan lett havne i den uhyggelige dalen. KI-algoritmer trenes på enorme datasett og lærer statistiske mønstre for hvordan ting "vanligvis" ser ut basert på tekstbeskrivelser (How Does AI Image Generation Work?, 2024). Selv om de kan generere imponerende resultater, kan de introdusere subtile feil – en unaturlig hudtekstur, feil antall fingre, en rar skygge, et "dødt" blikk – som vi mennesker, med vår fininnstilte persepsjon for artsfrender, raskt oppdager. Disse små avvikene kan trigge ubehaget assosiert med den uhyggelige dalen, og potensielt føre til en mer negativ emosjonell opplevelse sammenlignet med menneskeskapte bilder.'),
        ('h3', '1.2.5 Ytterligere teoretiske perspektiver'),
        ('p', 'Utover Uncanny Valley, kan andre teorier bidra til å forklare emosjonelle responser på KI-kunst:'),
        ('list', '* Nyhetseffekt (Novelty Effect): Mennesker har en tendens til å rette oppmerksomhet mot og vise økt fysiologisk aktivering (arousal) i møte med nye eller uventede stimuli. KI-generert kunst er for mange fortsatt nytt, noe som kan føre til sterkere innledende reaksjoner, enten positive (interesse, fascinasjon) eller negative (forvirring, skepsis), som kan påvirke skårer på intensitet (Spm 2).'),
        ('list', '* Prosesseringflyt (Processing Fluency): Teorien postulerer at jo lettere det er for hjernen å prosessere en stimulus, desto mer positivt vurderes den ofte (Reber, Schwarz, & Winkielman, 2004). KI-bilder, trent på enorme datasett, kan noen ganger destillere essensen av et konsept på en måte som gjør dem visuelt "enkle" eller prototypiske, til tross for potensiell overflatekompleksitet. Dette *kan* føre til høyere skårer på behagelighet (Spm 3) hvis de oppleves som lette å "forstå" visuelt.'),
        ('list', '* Aktiveringsteori og estetikk (f.eks. Berlyne): Daniel Berlyne knyttet estetisk preferanse til nivået av fysiologisk aktivering (arousal) en stimulus vekker. Faktorer som nyhet, kompleksitet, usikkerhet og overraskelse ("collative properties") påvirker aktiveringsnivået. Det antas ofte at et moderat nivå av aktivering oppleves som mest behagelig. KI-bilder kan ha høy grad av nyhet og kanskje kompleksitet, noe som kan føre til høy aktivering, som igjen kan påvirke både intensitets- og behagelighetsvurderinger.'),
        ('list', '* Teknologisk Sublimé: Begrepet beskriver følelsen av ærefrykt, undring, men kanskje også uro eller overveldelse, som kan oppstå i møte med kraftfull og avansert teknologi (Nye, 1994). KIs evne til å skape kunst kan utløse slike følelser, noe som kan bidra til høyere intensitetsskårer.'),
        ('h2', '1.3 Metode'),
        ('p', 'For å undersøke om menneskets emosjonelle respons varierer avhengig av om et bilde er skapt av kunstig intelligens eller et menneske, benyttet vi en eksperimentell metode med et innen-gruppe design (hver deltaker ble eksponert for begge betingelser: KI og menneskeskapt) og mellom-gruppe sammenligning (aldersgruppene). Ved å systematisk presentere par av KI-genererte og menneskeskapte bilder med lignende tematikk, og deretter måle deltakernes selvrapporterte emosjonelle responser, søkte vi å avdekke eventuelle signifikante forskjeller knyttet til bildenes opprinnelse og deltakernes alder.'),
        ('h2', '1.4 Variabler'), ('h3', '1.4.1 Uavhengige variabler'),
        ('p', 'De uavhengige variablene er de faktorene vi manipulerte eller kontrollerte for å observere effekten på den avhengige variabelen. I dette eksperimentet hadde vi to uavhengige variabler:'),
        ('list_num', '1.  Bildets Opprinnelse: Denne variabelen har to nivåer:\n    * KI-generert (bilde skapt av en KI-modell)\n    * Menneskeskapt (bilde skapt av en menneskelig kunstner)'),
        ('list_num', '2.  Aldersgruppe: Denne variabelen har to nivåer:\n    * Yngre gruppe (16-21 år)\n    * Eldre gruppe (30-50 år)'),
        ('p', 'Vi valgte disse variablene for å undersøke den primære problemstillingen om KI vs. menneskeskapt kunst, og for å utforske om alder, og potensielt ulik grad av digital kjennskap, modulerer denne responsen.'),
        ('h3', '1.4.2 Avhengige variabler'),
        ('p', 'Den avhengige variabelen er det vi måler for å se effekten av de uavhengige variablene. I dette studiet er den avhengige variabelen deltakernes emosjonelle respons, som ble operasjonalisert og målt gjennom tre selvrapporteringsspørsmål stilt under observasjonen av hvert bilde:'),
        ('list_num', '1.  Opplevd emosjon: (Åpent spørsmål: "Hvilken emosjon føler du?") - Gir kvalitativ data om typen følelse.'),
        ('list_num', '2.  Emosjonell intensitet: (Lukket skala 1-10: "Hvor sterkt reagerer du?") - Gir kvantitativ data om styrken på følelsen.'),
        ('list_num', '3.  Behagelighet (Valens): (Lukket skala 1-10: "Hvor behagelig opplever du bildet?") - Gir kvantitativ data om den positive/negative vurderingen av opplevelsen.'),
        ('h2', '1.5 Hypoteser'),
        ('p', 'Basert på det teoretiske rammeverket, spesielt teorien om den uhyggelige dalen og antakelser om ulik digital eksponering, formulerte vi følgende hypoteser før datainnsamlingen:'),
        ('list', '* H1 (Intensitet): Deltakerne vil rapportere *lavere* emosjonell intensitet (lavere score på Spm 2) for KI-genererte bilder sammenlignet med menneskeskapte bilder, på grunn av potensiell distanse eller mangel på "menneskelig touch".'),
        ('list', '* H2 (Behagelighet): Deltakerne vil rapportere *lavere* behagelighet (lavere score på Spm 3) for KI-genererte bilder sammenlignet med menneskeskapte bilder, spesielt hvis KI-bildene trigger en "uhyggelig dal"-effekt.'),
        ('list', '* H3 (Alderseffekt): Den eldre aldersgruppen (30-50) vil vise en *større* negativ forskjell i respons (både intensitet og behagelighet) mot KI-bilder sammenlignet med den yngre gruppen (16-21), som antas å være mer vant til digitalt og KI-generert innhold.'),
        ('list', '* H4 (Emosjonsgjenkjenning): Deltakerne vil i mindre grad rapportere en emosjon som samsvarer med bildets intenderte tema (målt ved "CorrectEmotion" i analysen) for KI-bilder sammenlignet med menneskeskapte bilder.'),
        ('page_break', ''),
        ('h1', '2.0 Forsøksdesign'), ('h2', '2.1 Utvalg'),
        ('p', 'Utvalget besto av totalt 16 deltakere, rekruttert gjennom bekvemmelighetsutvalg (f.eks. via sosiale nettverk, bekjente). Deltakerne ble fordelt likt i to aldersgrupper:'),
        ('list', '* Gruppe 1: 8 deltakere i alderen 16-21 år (4 kvinner, 4 menn). Gjennomsnittsalder: 18.4 år.'),
        ('list', '* Gruppe 2: 8 deltakere i alderen 30-50 år (4 kvinner, 4 menn). Gjennomsnittsalder: 40.6 år.'),
        ('p', 'Alle deltakere ga informert muntlig samtykke før deltakelse. De ble informert om studiens generelle formål (å undersøke emosjonell respons på bilder), men ble *ikke* eksplisitt fortalt at noen av bildene var KI-genererte. Dette ble gjort for å unngå at forventninger knyttet til KI skulle påvirke deres umiddelbare responser (en form for blinding). Deltakerne ble forsikret om anonymitet, og kun informasjon om alder og kjønn ble registrert. De ble debriefet om studiens fulle formål, inkludert bruken av KI-bilder, etter eksperimentets avslutning.'),
        ('h2', '2.2 Valg av bilder'),
        ('p', 'Stimulusmaterialet besto av 8 bildepar, totalt 16 bilder. Hvert par inneholdt:'),
        ('list_num', '1.  Ett menneskeskapt kunstverk.'),
        ('list_num', '2.  Ett KI-generert bilde.'),
        ('p', 'Bildene innenfor hvert par ble valgt eller generert for å representere den *samme* intenderte emosjonen eller temaet. Vi baserte valget av emosjoner på Paul Ekmans seks grunnleggende emosjoner (glede, tristhet, sinne, frykt, avsky, overraskelse), og utvidet med to tilleggsfølelser som ofte er relevante i kunstkontekst: nostalgi og kjedsomhet.'),
        ('p', 'De 8 menneskeskapte bildene ble valgt fra nettressursen "Art Emotions Map" (Art Emotions Map, u.å.), basert på deres tilknyttede emosjonelle merkelapper. Deretter brukte vi KI-modellen "Gemini" (via Googles verktøy) til å generere 8 korresponderende bilder. Vi ga KI-modellen tekstbeskrivelser ("prompts") som siktet mot å fremkalle de samme 8 emosjonene/temaene (kjedsomhet, nostalgi, frykt, tristhet, glede, avsky, overraskelse, sinne).'),
        ('p', 'Bildeparene ble deretter satt sammen basert på deres felles emosjonelle tema. Målet var å ha to bilder per tema, ett fra hver opprinnelseskategori (menneske/KI), for direkte sammenligning av responsene.'),
        ('h2', '2.3 Eksperimentets gjennomføring'),
        ('p', 'Eksperimentet ble gjennomført individuelt med hver deltaker i rolige omgivelser for å minimere distraksjoner. Prosedyren var som følger:'),
        ('list_num', '1.  Instruksjoner: Deltakeren fikk standardiserte instruksjoner om oppgaven: de skulle se på en serie bilder og svare på noen spørsmål om følelsene bildene vekket.'),
        ('list_num', '2.  Bildevisning og datainnsamling:\n    * For hvert av de 8 emosjonelle temaene ble deltakeren vist *enten* det menneskeskapte *eller* det KI-genererte bildet først. Rekkefølgen (KI først vs. Menneske først) ble forsøkt variert mellom deltakere for å motvirke systematisk rekkefølgeeffekt, selv om dette ikke var strengt kontrabalansert i dette pilotstudiet.\n    * Mens bildet var synlig på skjermen, stilte eksperimentleder de tre standardiserte spørsmålene (se 2.4). Svarene ble notert.\n    * Deretter ble det *andre* bildet i paret (med samme tema, men motsatt opprinnelse) vist.\n    * De samme tre spørsmålene ble stilt igjen, og svarene notert.'),
        ('list_num', '3.  Gjentakelse: Denne prosessen (vise bilde 1, spørre; vise bilde 2, spørre) ble gjentatt for alle 8 bildeparene/emosjonelle temaene.'),
        ('list_num', '4.  Debriefing: Etter siste bilde ble deltakeren takket for deltakelsen og informert om studiens fulle hensikt, inkludert skillet mellom KI-genererte og menneskeskapte bilder. Eventuelle spørsmål ble besvart.'),
        ('p','Varigheten per deltaker var omtrent 15-20 minutter.'),
        ('h2', '2.4 Innsamling av data'),
        ('p', 'Datainnsamlingen skjedde via strukturert intervju med selvrapportering underveis i eksperimentet. For hvert av de 16 bildene stilte vi følgende tre spørsmål:'),
        ('list_num', '1.  Åpent spørsmål (Type emosjon): «Hvilken emosjon føler du når du ser dette bildet?»\n    * *Formål:* Å fange den primære subjektive følelsen i deltakerens egne ord (kvalitativ data).'),
        ('list_num', '2.  Lukket spørsmål (Intensitet): «På en skala fra 1 (ingen reaksjon) til 10 (veldig sterk reaksjon), hvor sterkt emosjonelt reagerer du på bildet?»\n    * *Formål:* Å kvantifisere den opplevde styrken eller intensiteten av den emosjonelle responsen (kvantitativ data).'),
        ('list_num', '3.  Lukket spørsmål (Behagelighet/Valens): «På en skala fra 1 (veldig ubehagelig) til 10 (veldig behagelig), hvor behagelig opplever du dette bildet?»\n    * *Formål:* Å kvantifisere den positive eller negative valensen av den emosjonelle opplevelsen (kvantitativ data).'),
        ('p','Disse spørsmålene ble valgt for å operasjonalisere den emosjonelle responsen langs tre dimensjoner: type, intensitet og valens/behagelighet. Svarene ble kodet og lagt inn i en CSV-fil (`eksperiment_data.csv`) for videre analyse med Python-skriptet.'),
        ('page_break', ''),
        ('h1', '3.0 Resultater og Dataanalyse'),
        ('p', 'Dataene fra de 16 deltakerne ble analysert ved hjelp av det medfølgende Python-skriptet. Analysen fokuserte på å beregne gjennomsnittsscorer for emosjonell intensitet (Spm 2) og behagelighet (Spm 3), samt andelen "korrekte" emosjonsidentifikasjoner (basert på Spm 1 vs. bildetema), fordelt på bildetype (Menneske vs. KI) og aldersgruppe (16-21 vs. 30-50).'),
        ('h2', '3.1 Kvantitative resultater (Spm 2 & Spm 3)'),
        ('p', 'Gjennomsnittsscorene på skalaen 1-10 for intensitet (Spm 2) og behagelighet (Spm 3) var som følger:'),
        ('p_bold', 'Aldersgruppe 16-21 år:'),
        ('list', '* Intensitet (Spm 2): Menneskeskapte bilder: Gj.sn. 4.24 | KI-genererte bilder: Gj.sn. 6.50'),
        ('list', '* Behagelighet (Spm 3): Menneskeskapte bilder: Gj.sn. 4.00 | KI-genererte bilder: Gj.sn. 4.88'),
        ('p_bold', 'Aldersgruppe 30-50 år:'),
        ('list', '* Intensitet (Spm 2): Menneskeskapte bilder: Gj.sn. 5.50 | KI-genererte bilder: Gj.sn. 6.50'),
        ('list', '* Behagelighet (Spm 3): Menneskeskapte bilder: Gj.sn. 4.25 | KI-genererte bilder: Gj.sn. 4.50'),
        ('p_bold', 'Oppsummering av trender:'),
        ('list', '* Intensitet: Begge aldersgrupper rapporterte *høyere* gjennomsnittlig emosjonell intensitet for KI-bildene enn for de menneskeskapte bildene. Forskjellen var størst i den yngre gruppen.'),
        ('list', '* Behagelighet: Begge aldersgrupper rapporterte en marginalt *høyere* gjennomsnittlig behagelighet for KI-bildene sammenlignet med de menneskeskapte bildene. Igjen var forskjellen noe større i den yngre gruppen.'),
        ('p', 'Disse resultatene er visualisert i Graf 1 (under), Graf 2-5 (Appendix 7.1).'),
        ('img', 'all_spm'),
        ('h2', '3.2 Kvalitative resultater (Spm 1 - "Rett Emosjon")'),
        ('p', 'Analysen inkluderte en vurdering av om den emosjonen deltakeren rapporterte i Spm 1 samsvarte med det *intenderte* emosjonelle temaet for bildet (f.eks. om et bilde ment å vise "frykt" faktisk fremkalte en respons som "frykt"). Dette ble kodet som "CorrectEmotion" (1 for samsvar, 0 for ikke samsvar) basert på forhåndsdefinerte regler (se Appendix 7.3).'),
        ('p', 'Andelen svar som ble vurdert som "Rett Emosjon" var:'),
        ('p_bold', 'Aldersgruppe 16-21 år:'),
        ('list', '* Menneskeskapte bilder: 25.0% rett emosjon.'),
        ('list', '* KI-genererte bilder: 50.0% rett emosjon.'),
        ('p_bold', 'Aldersgruppe 30-50 år:'),
        ('list', '* Menneskeskapte bilder: 25.0% rett emosjon.'),
        ('list', '* KI-genererte bilder: 25.0% rett emosjon.'),
        ('p_bold', 'Oppsummering av trender:'),
        ('list', '* Den generelle andelen "rett emosjon" var relativt lav på tvers av alle betingelser, noe som indikerer at deltakernes subjektive opplevelse ofte avvek fra bildets intenderte tema.'),
        ('list', '* I den yngre gruppen (16-21) var andelen "rett emosjon" markant høyere for KI-bildene (50%) enn for de menneskeskapte bildene (25%).'),
        ('list', '* I den eldre gruppen (30-50) var det ingen forskjell i andelen "rett emosjon" mellom KI-bilder og menneskeskapte bilder (begge 25%).'),
        ('p', 'Dette er visualisert i Graf 6 (under og Appendix 7.1). En detaljert oversikt over hvert svar finnes i Graf 7 (Tabell) i Appendix 7.2.'),
        ('img', 'emo_match'),
        ('page_break', ''),
        ('h1', '4.0 Validitet og reliabilitet'),
        ('p', 'Vurdering av studiens validitet (måler vi det vi tror vi måler?) og reliabilitet (er målingene stabile og konsistente?) er essensielt for å tolke funnene.'),
        ('p_bold', 'Validitet:'),
        ('list', '* Indre validitet: Handler om hvorvidt vi kan trekke sikre konklusjoner om årsakssammenhenger (dvs. at forskjeller i respons *skyldes* bildets opprinnelse eller aldersgruppe). Styrker her er bruken av et eksperimentelt design med sammenligning av KI og menneskeskapte bilder med samme tema. Svakheter inkluderer potensielle konfunderende variabler, som subtile forskjeller i bildekvalitet eller stil *utover* opprinnelsen, og rekkefølgeeffekter (nevnt i 4.1.2). Blindingen av deltakere for bildets opprinnelse styrker den indre validiteten ved å redusere forventningseffekter.'),
        ('list', '* Ytre validitet (Generaliserbarhet): Handler om i hvilken grad resultatene kan generaliseres til andre personer, settinger og bilder. Denne er begrenset på grunn av et lite utvalg (N=16) og ikke-tilfeldig utvalgsmetode (bekvemmelighetsutvalg). Funnene er primært gyldige for *disse* deltakerne og *disse* bildene. Generalisering til befolkningen generelt, eller til andre typer kunst/bilder, må gjøres med stor forsiktighet.'),
        ('list', '* Begrepsvaliditet (Construct Validity): Handler om hvorvidt våre operasjonaliseringer (spørsmålene) faktisk måler de teoretiske begrepene (emosjonell respons). Selvrapportering av intensitet og behagelighet er vanlige metoder, men subjektive. Åpne spørsmål om emosjon gir rikere data, men "rett emosjon"-analysen avhenger av både bildets evne til å formidle *og* de forhåndsdefinerte reglene for samsvar, noe som kan være en svakhet. Spørsmålet er om skalaene og det åpne spørsmålet fanger hele kompleksiteten i en estetisk emosjonell opplevelse.'),
        ('p_bold', 'Reliabilitet:'),
        ('list', '* Reliabiliteten til selvrapporteringsskalaene (Spm 2 og 3) antas å være akseptabel for denne type måling, men ble ikke formelt testet (f.eks. med test-retest).'),
        ('list', '* Konsistensen i prosedyrer og instruksjoner var viktig for å sikre at alle deltakere ble behandlet likt, noe som styrker reliabiliteten.'),
        ('list', '* Kodingen av det åpne spørsmålet (Spm 1) for "rett emosjon"-analysen ble gjort basert på klare regler i skriptet, noe som sikrer konsistens i *denne* analysen, men selve reglene kan diskuteres (se over).'),
        ('h2', '4.1 Styrker og svakheter'), ('h3', '4.1.1 Styrker'),
        ('list', '* Eksperimentelt design: Tillater en viss grad av kausal inferens ved direkte sammenligning av responser på KI vs. menneskeskapte bilder innenfor samme tema.'),
        ('list', '* Bruk av bildepar: Kontrollerer for det emosjonelle temaet når man sammenligner KI og menneskeskapt opprinnelse.'),
        ('list', '* Kombinasjon av mål: Bruk av både kvantitative skalaer (intensitet, behagelighet) og kvalitative data (type emosjon) gir et mer nyansert bilde av responsen.'),
        ('list', '* Inkludering av aldersgrupper: Muliggjør utforsking av potensielle aldersrelaterte forskjeller i persepsjon av KI-kunst.'),
        ('list', '* Blinding: Deltakerne visste ikke på forhånd hvilke bilder som var KI-generert, noe som reduserer bias.'),
        ('h3', '4.1.2 Svakheter'),
        ('list', '* Lite og ikke-representativt utvalg: Som nevnt, N=16 og bekvemmelighetsutvalg begrenser generaliserbarheten sterkt. Resultatene kan være tilfeldige eller spesifikke for denne gruppen.'),
        ('list', '* Potensielle rekkefølgeeffekter: Selv om rekkefølgen ble variert, var den ikke strengt kontrabalansert. Responsen på det andre bildet i et par kan ha blitt påvirket av å ha sett det første.'),
        ('list', '* Subjektivitet i selvrapportering: Emosjonelle responser er personlige og kan påvirkes av dagsform, tidligere erfaringer og tolkning av skalaene.'),
        ('list', '* Begrenset operasjonalisering: Målingene fanger kanskje ikke alle aspekter ved en kompleks estetisk opplevelse. Fysiologiske mål (puls, hudledningsevne) kunne gitt tilleggsinformasjon.'),
        ('list', '* Kvalitet og valg av KI-bilder: Resultatene kan være spesifikke for *disse* KI-bildene generert av "Gemini". Andre KI-modeller eller bilder med tydeligere "uncanny valley"-trekk kunne gitt andre resultater. Likeledes kan kvaliteten og appellen til de valgte menneskeskapte bildene ha påvirket sammenligningen.'),
        ('list', '* Definisjon av "Rett Emosjon": Vurderingen av om en opplevd emosjon var "rett" er basert på en forenklet matching mot bildetemaet og forhåndsdefinerte synonymer. Kunst er ofte ambiguøs, og ulike tolkninger kan være like valide. Den lave treffprosenten kan reflektere dette mer enn en "feil" hos deltakeren eller bildet.'),
        ('list', '* Etiske betraktninger: Bruk av blinding (ikke informere om KI før etterpå) er vanlig i forskning for å unngå bias, men krever grundig debriefing.'),

        ('h2', '4.2 Drøfting'), # (Utvidet drøfting fra forrige svar)
        ('p', 'Resultatene fra dette piloteksperimentet presenterer et interessant, og til dels kontraintuitivt, bilde av hvordan unge og middelaldrende voksne responderer emosjonelt på KI-generert versus menneskeskapt kunst i denne spesifikke konteksten. Funnene utfordrer våre innledende hypoteser og inviterer til en bredere teoretisk diskusjon.'),
        ('p_bold', 'Hovedfunn og avvik fra hypoteser:'),
        ('p', 'De mest slående funnene var at KI-genererte bilder, på tvers av begge aldersgrupper, i gjennomsnitt ble vurdert til å fremkalle *høyere* emosjonell intensitet (Spm 2) og marginalt *høyere* behagelighet (Spm 3) enn de menneskeskapte bildene. Dette står i direkte kontrast til hypotesene (H1 og H2) som, basert primært på Uncanny Valley-teorien, forutså det motsatte – at KI-bilder ville oppleves som mindre intense og mindre behagelige. Videre ble hypotesen om at den eldre gruppen ville reagere mest negativt på KI (H3) heller ikke støttet; tvert imot så vi de største forskjellene (i favør av KI) i den yngre gruppen. Også hypotesen om dårligere emosjonsgjenkjenning for KI (H4) ga blandede resultater, med *bedre* gjenkjenning for KI i den yngre gruppen.'),
        ('p_bold', 'Hvorfor sviktet Uncanny Valley-hypotesen her?'),
        ('p', 'Fraværet av en tydelig negativ reaksjon på KI-bildene, og til og med en tendens til det motsatte, antyder at Uncanny Valley-effekten enten ikke ble utløst av disse spesifikke bildene, eller at andre faktorer veide tyngre. Flere årsaker er mulige:'),
        ('list', '* Bildematerialet: KI-bildene generert av Gemini for dette forsøket unngikk kanskje de groveste feilene (som feil antall fingre, forvrengte ansikter) som ofte forbindes med uhygge. De kan ha hatt en estetisk stil (f.eks. mer fargerik, glatt, teknisk "perfekt") som appellerte til deltakerne, uavhengig av subtile KI-markører. Samtidig kan de valgte menneskeskapte bildene ha vært mer ambivalente eller mindre umiddelbart engasjerende.'),
        ('list', '* Kontekst: Uncanny Valley er kanskje mest potent for representasjoner som *nesten* er fotorealistiske mennesker (roboter, CGI-ansikter). Vårt utvalg inkluderte mer varierte motiver og stiler, hvor forventningen om perfekt menneskelikhet var lavere.'),
        ('list', '* Målemetode: Selv om et bilde kan inneholde elementer som teoretisk sett kan trigge uhygge, er det ikke sikkert dette fanges opp av en enkel 1-10 skala for "behagelighet". Følelsen kan være mer kompleks eller blandet.'),
        ('p_bold', 'Alternative teoretiske forklaringer på funnene:'),
        ('p', 'Gitt at KI-bildene skåret *høyere* på intensitet og likt/høyere på behagelighet, må vi vurdere andre teorier introdusert i seksjon 1.2.5:'),
        ('list', '* Nyhetseffekt og Teknologisk Sublimé: KI-kunst er et relativt nytt fenomen. Den økte intensiteten kan reflektere en sterkere orienteringsrespons, fascinasjon, eller ærefrykt ("wow-effekt") knyttet til det nye og teknologisk avanserte. Dette kan være spesielt fremtredende hos den yngre gruppen (16-21) som viste størst forskjell i intensitet. De er kanskje mer opptatt av teknologiske fremskritt, eller de opplever en sterkere kognitiv eller emosjonell respons når de konfronteres med noe som utfordrer etablerte kategorier (hva er "kunst"?).'),
        ('list', '* Aktiveringsteori (Berlyne): KI-bildene kan ha hatt høyere "kollative egenskaper" – mer nyhet, kanskje opplevd kompleksitet eller overraskelse – som førte til et høyere nivå av fysiologisk aktivering (arousal). Denne økte aktiveringen kan ha blitt tolket og rapportert som høyere emosjonell "intensitet". At behageligheten også var marginalt høyere, kan tyde på at dette økte aktiveringsnivået (for disse bildene og deltakerne) ikke krysset grensen til å bli aversivt, men snarere ble opplevd som engasjerende eller interessant.'),
        ('list', '* Prosesseringflyt: Den marginalt høyere behageligheten, og spesielt den *bedre* "korrekte" emosjonsgjenkjenningen for KI-bilder hos den yngre gruppen, *kan* spekulativt knyttes til prosesseringsflyt. Hvis KI-bildene, trent på massive datasett, tenderer mot å produsere mer arketypiske eller visuelt destillerte representasjoner av en emosjon (f.eks. en "gjennomsnittlig" fremstilling av glede), kan disse være kognitivt enklere å bearbeide og kategorisere enn et unikt, kanskje mer subtilt eller tvetydig, menneskeskapt verk. Lettere prosessering korrelerer ofte med økt preferanse og en følelse av "riktighet". Dette er dog en svært spekulativ tolkning som krever mer forskning.'),
        ('p_bold', 'Aldersforskjeller:'),
        ('p', 'At den yngre gruppen reagerte sterkere (høyere intensitet for KI) og mer "korrekt" (bedre emosjonsgjenkjenning for KI) enn den eldre gruppen, kan reflektere generasjonsforskjeller i eksponering og holdninger til digitalt og KI-generert innhold. Yngre individer har vokst opp med digital teknologi og kan ha andre estetiske preferanser eller skjemaer for å tolke slike bilder. De kan være mer åpne for, eller fascinerte av, KI-estetikk, mens den eldre gruppen kanskje har mer etablerte kunstpreferanser eller er mer skeptiske. Den eldre gruppens lavere og mer like responser på tvers av bildetyper kan tyde på en mer avmålt eller mindre differensiert reaksjon i denne konteksten.'),
        ('p_bold', 'Begrensninger og veien videre:'),
        ('p', 'Det er avgjørende å understreke studiens begrensninger, primært det svært lille og ikke-representative utvalget. Funnene kan ikke generaliseres og må betraktes som foreløpige indikasjoner som krever replikering og utdyping. Definisjonen av "rett emosjon" er også en betydelig forenkling av den komplekse prosessen det er å tolke kunst.'),
        ('p', 'Fremtidig forskning bør adressere disse svakhetene ved å bruke større, mer mangfoldige utvalg og strengere kontrollerte stimuli (f.eks. bedre matching av KI og menneskelige bilder på stil og kompleksitet, eller testing av bilder spesifikt designet for å ligge i Uncanny Valley). Mer sofistikerte målemetoder, inkludert fysiologiske mål (hjerteratevariabilitet, hudkonduktans) og kvalitative dybdeintervjuer, kan gi rikere innsikt i de underliggende prosessene. Det vil også være viktig å undersøke effekten av å *informere* deltakerne om bildets opprinnelse på forhånd – hvordan påvirker kunnskapen om at noe er KI-skapt vår opplevelse?'),
        ('p', 'Denne studien, til tross for sine begrensninger, antyder at forholdet mellom mennesker og KI-generert kunst er mer nyansert enn en enkel frykt for det "uhyggelige". Faktorer som nyhet, teknologisk fascinasjon, prosesseringsflyt og individuelle/generasjonsmessige forskjeller ser ut til å spille en vesentlig rolle, og fortjener grundigere utforskning.'),
        ('page_break', ''),

        ('h1', '5.0 Konklusjon'),
        ('p', 'Dette piloteksperimentet undersøkte emosjonelle responser (intensitet, behagelighet, type følelse) til KI-genererte versus menneskeskapte bilder hos en yngre (16-21) og en eldre (30-50) aldersgruppe. I motsetning til våre hypoteser basert på "Uncanny Valley"-teorien, fant vi at deltakerne i begge aldersgrupper rapporterte høyere emosjonell intensitet og marginalt høyere behagelighet for KI-bildene sammenlignet med de menneskeskapte bildene i vårt utvalg. Den yngre gruppen viste de største forskjellene i favør av KI-bilder, og hadde også en høyere andel "korrekt" emosjonsidentifikasjon for KI-bilder enn for menneskeskapte bilder.'),
        ('p', 'Selv om studiens generaliserbarhet er begrenset av utvalgsstørrelsen og -metoden, antyder funnene at KI-generert kunst kan være emosjonelt engasjerende og ikke nødvendigvis oppleves som mindreverdig eller uhyggelig sammenlignet med menneskeskapt kunst, spesielt for yngre publikum. Resultatene understreker kompleksiteten i menneskers interaksjon med KI-skapt innhold og peker på behovet for videre forskning for å forstå de psykologiske mekanismene som ligger til grunn for våre responser på denne nye formen for visuelt uttrykk.'),
        ('page_break', ''),
        ('h1', '6.0 Referanser'),
        ('ref', 'Art Emotions Map. (u.å.). Hentet 31. mars 2025, fra https://artsexperiments.withgoogle.com/art-emotions-map/'),
        ('ref', 'Berlyne, D. E. (1971). *Aesthetics and psychobiology*. Appleton-Century-Crofts.'),
        ('ref', 'Christensen, B. (2014, mars 8). Hjernen speiler andres oppførsel. *forskning.no*. https://www.forskning.no/psykologi-menneskekroppen/hjernen-speiler-andres-oppforsel/576610'),
        ('ref', 'Den uhyggelige dal. (2024). I *Wikipedia*. Hentet 31. mars 2025, fra https://no.wikipedia.org/w/index.php?title=Den_uhyggelige_dal&oldid=24845205'),
        ('ref', 'Derfor er robotter og klovne så uhyggelige. (2020, januar 16). *Videnskab.dk*. https://videnskab.dk/kultur-samfund/derfor-er-robotter-og-klovne-saa-uhyggelige/'),
        ('ref', 'Dr. Paul Ekman. (u.å.). Paul Ekman Group. Hentet 31. mars 2025, fra https://www.paulekman.com/about/paul-ekman/'),
        ('ref', 'Emotion classification. (2024). In *Wikipedia*. Hentet 31. mars 2025, fra https://en.wikipedia.org/w/index.php?title=Emotion_classification&oldid=1261117450'),
        ('ref', 'How Does AI Image Generation Work? - Hypotenuse AI. (2024, mars 13). Hentet 31. mars 2025, fra https://www.hypotenuse.ai/blog/how-do-ai-image-generators-work'),
        ('ref', 'MacDorman, K. F. (2005). *Androids as an experimental apparatus: Why is there an uncanny valley and can we exploit it?* [Grafikk hentet fra Wikimedia Commons]. CogSci-2005 Workshop: Toward Social Mechanisms of Android Science. https://commons.wikimedia.org/wiki/File:MoriUncannyValley.svg'),
        ('ref', 'Nye, D. E. (1994). *American Technological Sublime*. MIT Press.'),
        ('ref', 'Oct. 13, K. M. on, 2009, 11. a.m. (u.å.). Like humans, monkeys fall into the «uncanny valley». *Princeton University News*. Hentet 31. mars 2025, fra https://www.princeton.edu/news/2009/10/13/humans-monkeys-fall-uncanny-valley'),
        ('ref', 'Reber, R., Schwarz, N., & Winkielman, P. (2004). Processing fluency and aesthetic pleasure: Is beauty in the perceiver’s processing experience? *Personality and Social Psychology Review, 8*(4), 364–382. https://doi.org/10.1207/s15327957pspr0804_3'),
        ('ref', 'Silvia, P. J. (2005). Emotional Responses to Art: From Collation and Arousal to Cognition and Emotion. *Review of General Psychology, 9*(4), 342–357. https://doi.org/10.1037/1089-2680.9.4.342'),
        ('ref', 'Speilnevroner. (2024). I *Wikipedia*. Hentet 31. mars 2025, fra https://no.wikipedia.org/w/index.php?title=Speilnevroner&oldid=24468535'),
        ('ref', 'Steckenfinger, S. A., & Ghazanfar, A. A. (2009). Monkey visual behavior falls into the uncanny valley. *Proceedings of the National Academy of Sciences, 106*(43), 18392–18396. https://doi.org/10.1073/pnas.0910063106'),
        ('ref', 'Svartdal, F. (2025). Emosjon. I *Store norske leksikon*. Hentet 31. mars 2025, fra https://snl.no/emosjon'),
        ('ref', 'Yin, P.-Y. (Ed.). (2008). *Pattern Recognition: Techniques, Technology and Applications*. I-Tech Education and Publishing.'),
        ('page_break', ''),

        ('h1', '7.0 Appendix'), ('h2', '7.1 Genererte Grafer'),
        ('p', 'Følgende grafer ble generert basert på data fra eksperimentet og er referert til i Resultat- og Drøftingsseksjonene. (Bildene settes inn nedenfor).'),
        ('list_num', '1.  (Graf 1: `all_spm`) `graf_alle_grupper_alle_spm...png`: Sammenligning Intensitet/Behagelighet for begge grupper.'), ('img', 'all_spm'),
        ('list_num', '2.  (Graf 2: `g1_spm`) `graf_gruppe_16-21...png`: Resultater for gruppe 16-21.'), ('img', 'g1_spm'),
        ('list_num', '3.  (Graf 3: `g2_spm`) `graf_gruppe_30-50...png`: Resultater for gruppe 30-50.'), ('img', 'g2_spm'),
        ('list_num', '4.  (Graf 4: `spm2_comp`) `graf_alle_grupper_spm2...png`: Sammenligning Intensitet (Spm 2).'), ('img', 'spm2_comp'),
        ('list_num', '5.  (Graf 5: `spm3_comp`) `graf_alle_grupper_spm3...png`: Sammenligning Behagelighet (Spm 3).'), ('img', 'spm3_comp'),
        ('list_num', '6.  (Graf 6: `emo_match`) `graf_rett_emosjon_prosent...png`: Prosentandel "Rett Emosjon".'), ('img', 'emo_match'),
        ('h2', '7.2 Tabell: Opplevd Emosjon Status'),
        ('p', 'Tabellen nedenfor (generert som `tabell_rett_emosjon_status...png`, Graf 7) viser hver enkelt respons på Spm 1.'),
        ('img', 'emo_table'),
        ('h2', '7.3 Definisjon av "Rett Emosjon"'),
        ('p', 'For analysen av Spm 1 ("Hvilken emosjon føler du?") ble følgende regler brukt for å avgjøre om en deltagers svar ble kodet som "Rett Emosjon":'),
        ('list', '* Kjedsomhet: [\'kjedsomhet\', \'kjedelig\']'),
        ('list', '* Nostalgi: [\'nostalgi\', \'nostalgisk\']'),
        ('list', '* Sinne: [\'sinne\']'),
        ('list', '* Frykt: [\'frykt\', \'krig/frykt\']'),
        ('list', '* Avsky: [\'avsky\', \'kvalme\']'),
        ('list', '* Glede: [\'glede\', \'glad\', \'håp\']'),
        ('list', '* Overraskelse: [\'overraskelse\', \'overrasket\', \'overveldende\']'),
        ('list', '* Tristhet: [\'tristhet\', \'trist\']'),
        ('p', 'Alle andre svar for et gitt bildetema ble kodet som "Ikke rett".')
    ]
    return report_content

# ===============================================
# 3. HOVEDPROGRAM: KJØR ANALYSE, BYGG DOKUMENT
#    (Identisk med forrige versjon)
# ===============================================

initial_data = { # Data inkludert her for fullstendighet
    'AgeGroup': ['16-21', '16-21', '16-21', '16-21', '16-21', '16-21', '16-21', '16-21','16-21', '16-21', '16-21', '16-21', '16-21', '16-21', '16-21', '16-21','30-50', '30-50', '30-50', '30-50', '30-50', '30-50', '30-50', '30-50','30-50', '30-50', '30-50', '30-50', '30-50', '30-50', '30-50', '30-50'],
    'ParticipantID': ['Jente 18a', 'Jente 18a', 'Gutt 18a', 'Gutt 18a', 'Gutt 17', 'Gutt 17', 'Jente 18b', 'Jente 18b','Gutt 19', 'Gutt 19', 'Gutt 18b', 'Gutt 18b', 'Jente 19', 'Jente 19', 'Jente 20', 'Jente 20','Kvinne 36', 'Kvinne 36', 'Mann 38', 'Mann 38', 'Mann 50', 'Mann 50', 'Mann 44', 'Mann 44','Kvinne 47', 'Kvinne 47', 'Mann 35', 'Mann 35', 'Kvinne 33', 'Kvinne 33', 'Kvinne 42', 'Kvinne 42'],
    'ImageType': ['Menneske', 'KI', 'Menneske', 'KI', 'Menneske', 'KI', 'Menneske', 'KI','Menneske', 'KI', 'Menneske', 'KI', 'Menneske', 'KI', 'Menneske', 'KI','Menneske', 'KI', 'Menneske', 'KI', 'Menneske', 'KI', 'Menneske', 'KI','Menneske', 'KI', 'Menneske', 'KI', 'Menneske', 'KI', 'Menneske', 'KI'],
    'ImageName': ['Kjedsomhet', 'Kjedsomhet', 'Nostalgi', 'Nostalgi', 'Frykt', 'Frykt', 'Tristhet', 'Tristhet','Glede', 'Glede', 'Avsky', 'Avsky', 'Overraskelse', 'Overraskelse', 'Sinne', 'Sinne','Kjedsomhet', 'Kjedsomhet', 'Sinne', 'Sinne', 'Nostalgi', 'Nostalgi', 'Frykt', 'Frykt','Avsky', 'Avsky', 'Glede', 'Glede', 'Overraskelse', 'Overraskelse', 'Tristhet', 'Tristhet'],
    'Spm1': ['kaos', 'Fred', 'ukomfortabel', 'Glede', 'sinne', 'håp', 'Tristhet', 'Tristhet','Nostalgi', 'Glede', 'Kaos', 'Kvalme', 'frykt', 'Overveldende', 'Sinne', 'Frykt','Sinne', 'Nostalgi', 'Krig/frykt', 'Frykt', 'nysgjerrighet', 'Avslappende', 'Frykt', 'Kaotisk','Forvirring', 'Avsky', 'Avslapning', 'Forvirring', 'Tristhet', 'Forvirring', 'Tristhet', 'Tristhet'],
    'Spm2': [3, 4, 6.9, 8, 6, 4, 3, 7,4, 6, 4, 8, 3, 7, 4, 8,5, 7, 2, 2, 8, 4, 5, 9,7, 10, 3, 4, 7, 7, 7, 9],
    'Spm3': [4, 8, 3, 8, 3, 8, 5, 3,5, 7, 3, 1, 4, 2, 5, 2,3, 9, 5, 5, 2, 4, 7, 2,2, 0, 7, 6, 2, 5, 6, 5]
}

generated_files = generate_plots_and_get_filenames(initial_data)

if CAN_RUN_DOCX and generated_files:
    print("\nStarter bygging av Word-dokument...")
    report_structure = get_full_report_text_complete() # Bruker den nye funksjonen med komplett tekst
    document = docx.Document()

    for item_type, content in report_structure:
        try:
            if item_type == 'title':
                p = document.add_paragraph(content, style='Title'); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif item_type == 'h1': document.add_heading(content, level=1)
            elif item_type == 'h2': document.add_heading(content, level=2)
            elif item_type == 'h3': document.add_heading(content, level=3)
            elif item_type == 'p': document.add_paragraph(content)
            elif item_type == 'p_bold': document.add_paragraph().add_run(content).bold = True
            elif item_type == 'p_center_bold':
                p = document.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; p.add_run(content).bold = True
            elif item_type == 'list':
                text = content.lstrip('*').strip(); document.add_paragraph(text, style='List Bullet')
            elif item_type == 'list_num':
                 text = content.split('.', 1)[-1].strip() if '.' in content else content
                 text = text.split(')', 1)[-1].strip() if ')' in text else text
                 while text and text[0].isdigit() or text.startswith('.'): text = text[1:].lstrip()
                 document.add_paragraph(text, style='List Number')
            elif item_type == 'ref':
                 p = document.add_paragraph(content); p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.first_line_indent = Inches(-0.25)
            elif item_type == 'toc': document.add_paragraph(content)
            elif item_type == 'toc_sub':
                 p = document.add_paragraph(content); p.paragraph_format.left_indent = Inches(0.25)
            elif item_type == 'toc_ssub':
                 p = document.add_paragraph(content); p.paragraph_format.left_indent = Inches(0.5)
            elif item_type == 'page_break': document.add_page_break()
            elif item_type == 'img':
                image_key = content
                if image_key in generated_files:
                    img_path = generated_files[image_key]
                    if os.path.exists(img_path):
                        try:
                            print(f"  Legger til bilde: {img_path}")
                            document.add_paragraph()
                            document.add_picture(img_path, width=Inches(6.0))
                            last_paragraph = document.paragraphs[-1]
                            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            document.add_paragraph()
                        except Exception as e:
                            print(f"  FEIL ved innsetting av bilde {img_path}: {e}")
                            document.add_paragraph(f"[FEIL: Kunne ikke sette inn {img_path} - {e}]")
                    else:
                        print(f"  FEIL: Bildfil ikke funnet på disk: {img_path}")
                        document.add_paragraph(f"[FEIL: Bildet {img_path} ble ikke funnet]")
                else:
                    print(f"  FEIL: Bildenøkkel '{image_key}' ikke funnet i genererte filer.")
                    document.add_paragraph(f"[FEIL: Bildedata for '{image_key}' mangler]")
            else: document.add_paragraph(str(content))
        except Exception as e:
             print(f"FEIL under prosessering av element: Type={item_type}, Feil={e}")
             try: document.add_paragraph(f"[FEIL i dokumentbygging: {e}]")
             except: pass

    save_timestamp = int(time.time())
    output_filename = f"rapport_KI_vs_Menneske_komplett_{save_timestamp}.docx"
    try:
        document.save(output_filename)
        print(f"\nWord-dokument forsøkt lagret som: {output_filename}")
        print("!!! VIKTIG: Denne filen er kun lagret internt i kjøremiljøet og kan IKKE hentes ut eller leveres til deg. !!!")
        print("\n(Simulerer sletting av midlertidige PNG-filer...)")
        # Opprydding kan aktiveres ved lokal kjøring
        # for filename in generated_files.values():
        #    if os.path.exists(filename): os.remove(filename)
    except Exception as e: print(f"\nFEIL ved lagring av Word-dokument: {e}")

elif not CAN_RUN_DOCX:
    print("\nWord-dokument ble ikke opprettet fordi 'python-docx' mangler.")
    if generated_files:
         print("Bilder ble generert, men ikke lagt inn i et dokument.")
         print("Genererte filnavn (kun internt):", list(generated_files.values()))
    else: print("Hverken bilder eller dokument ble generert (sannsynligvis datafeil).")
else: # generated_files was None
    print("\nKunne ikke generere bilder/data, Word-dokument ble ikke opprettet.")

print("\nProgram ferdig.")